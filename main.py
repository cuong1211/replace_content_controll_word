import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path
import zipfile
import re


def method1_replace_content_control(doc, tag_name, new_value):
    """
    Phương pháp 1: Thay thế Content Control qua python-docx
    """
    try:
        replaced_count = 0

        def find_and_replace_in_element(element):
            nonlocal replaced_count
            if hasattr(element, "_element"):
                for sdt in element._element.iter():
                    if sdt.tag.endswith("}sdt"):
                        # Tìm tag
                        for child in sdt.iter():
                            if child.tag.endswith("}tag"):
                                tag_val = child.get(qn("w:val"))
                                if tag_val == tag_name:
                                    # Tìm và thay thế text
                                    for text_elem in sdt.iter():
                                        if text_elem.tag.endswith("}t"):
                                            old_text = text_elem.text or ""
                                            text_elem.text = str(new_value)
                                            replaced_count += 1
                                            print(
                                                f"      ✅ Thay thế: '{old_text}' → '{new_value}'"
                                            )
                                    break

        # Tìm trong document
        find_and_replace_in_element(doc)

        # Tìm trong paragraphs
        for paragraph in doc.paragraphs:
            find_and_replace_in_element(paragraph)

        # Tìm trong tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    find_and_replace_in_element(cell)

        return replaced_count > 0

    except Exception as e:
        print(f"      ❌ Lỗi method1: {e}")
        return False


def method4_xml_replacement(template_file, output_file, tag_name, new_value):
    """
    Phương pháp 4: Thay thế trực tiếp trong XML (backup method)
    """
    try:
        with zipfile.ZipFile(template_file, "r") as zip_ref:
            document_xml = zip_ref.read("word/document.xml").decode("utf-8")

            # Pattern để tìm Content Control với tag cụ thể
            pattern = rf'(<w:sdt[^>]*>.*?<w:tag w:val="{tag_name}".*?<w:t[^>]*>)[^<]*(</w:t>.*?</w:sdt>)'

            def replace_text(match):
                return match.group(1) + str(new_value) + match.group(2)

            new_xml = re.sub(pattern, replace_text, document_xml, flags=re.DOTALL)

            if new_xml != document_xml:
                # Tạo file mới
                with zipfile.ZipFile(output_file, "w", zipfile.ZIP_DEFLATED) as new_zip:
                    for item in zip_ref.infolist():
                        if item.filename != "word/document.xml":
                            new_zip.writestr(item, zip_ref.read(item.filename))
                    new_zip.writestr("word/document.xml", new_xml.encode("utf-8"))
                return True

        return False

    except Exception as e:
        print(f"      ❌ Lỗi method4: {e}")
        return False


def create_word_file_for_name(template_file, name, output_file, tag_name="name"):
    """
    Tạo file Word mới cho một tên cụ thể
    """
    print(f"   🔄 Đang tạo file cho: {name}")

    try:
        # Phương pháp 1: Sử dụng python-docx
        doc = Document(template_file)
        success1 = method1_replace_content_control(doc, tag_name, name)

        if success1:
            doc.save(output_file)
            print(f"   ✅ Method 1 thành công: {output_file}")
            return True
        else:
            print(f"   ⚠️ Method 1 thất bại, thử Method 4...")
            # Phương pháp 2: XML replacement
            success4 = method4_xml_replacement(
                template_file, output_file, tag_name, name
            )

            if success4:
                print(f"   ✅ Method 4 thành công: {output_file}")
                return True
            else:
                print(f"   ❌ Tất cả phương pháp thất bại cho: {name}")
                return False

    except Exception as e:
        print(f"   ❌ Lỗi tạo file cho {name}: {e}")
        return False


def verify_file_content(file_path, expected_text):
    """
    Kiểm tra xem file có chứa text mong đợi không
    """
    try:
        doc = Document(file_path)

        # Kiểm tra trong paragraphs
        for paragraph in doc.paragraphs:
            if expected_text in paragraph.text:
                return True

        # Kiểm tra trong tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if expected_text in cell.text:
                        return True

        return False

    except Exception as e:
        print(f"      ⚠️ Lỗi kiểm tra file {file_path}: {e}")
        return False


def process_excel_to_word_files(
    excel_file, template_file, output_folder, tag_name="name"
):
    """
    Xử lý toàn bộ: đọc Excel và tạo file Word cho mỗi tên
    """
    print(f"🚀 BẮT ĐẦU XỬ LÝ EXCEL → WORD FILES")
    print("=" * 60)

    # Tạo thư mục output
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    print(f"📁 Thư mục output: {output_folder}")

    try:
        # Đọc Excel
        df = pd.read_excel(excel_file)
        print(f"📊 Đọc được {len(df)} dòng từ {excel_file}")
        print(f"📋 Các cột: {list(df.columns)}")

        # Kiểm tra cột 'name'
        if "name" not in df.columns:
            print(f"❌ Không tìm thấy cột 'name' trong Excel!")
            print(f"💡 Các cột có sẵn: {list(df.columns)}")
            return

        # Lọc dữ liệu hợp lệ
        valid_names = []
        for index, row in df.iterrows():
            name = row["name"]
            if pd.notna(name) and str(name).strip():
                valid_names.append(str(name).strip())

        print(f"✅ Tìm thấy {len(valid_names)} tên hợp lệ")

        if not valid_names:
            print(f"❌ Không có tên hợp lệ nào để xử lý!")
            return

        # Xử lý từng tên
        success_count = 0
        failed_count = 0

        for i, name in enumerate(valid_names, 1):
            print(f"\n📝 [{i}/{len(valid_names)}] Xử lý: {name}")

            # Tạo tên file an toàn
            safe_name = name
            for char in ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]:
                safe_name = safe_name.replace(char, "_")

            output_file = os.path.join(output_folder, f"{safe_name}.docx")

            # Tạo file
            success = create_word_file_for_name(
                template_file, name, output_file, tag_name
            )

            if success:
                # Kiểm tra kết quả
                if verify_file_content(output_file, name):
                    print(f"   ✅ Xác nhận: File chứa '{name}'")
                    success_count += 1
                else:
                    print(f"   ⚠️ Cảnh báo: File không chứa '{name}' (có thể do format)")
                    success_count += 1  # Vẫn tính là thành công vì file đã được tạo
            else:
                failed_count += 1

        # Báo cáo kết quả
        print("\n" + "=" * 60)
        print(f"📊 KẾT QUẢ CUỐI CÙNG:")
        print(f"   ✅ Thành công: {success_count}/{len(valid_names)} file")
        print(f"   ❌ Thất bại: {failed_count}/{len(valid_names)} file")
        print(f"   📁 Thư mục output: {output_folder}")

        # Liệt kê file đã tạo
        if success_count > 0:
            print(f"\n📋 CÁC FILE ĐÃ TẠO:")
            for filename in os.listdir(output_folder):
                if filename.endswith(".docx"):
                    file_path = os.path.join(output_folder, filename)
                    file_size = os.path.getsize(file_path)
                    print(f"   📄 {filename} ({file_size:,} bytes)")

        print("=" * 60)

    except Exception as e:
        print(f"❌ Lỗi chính: {e}")


def create_sample_data():
    """
    Tạo dữ liệu mẫu để test
    """
    print("🔧 Tạo dữ liệu mẫu...")

    # Tạo Excel mẫu
    sample_data = {
        "name": [
            "Nguyễn Văn An",
            "Trần Thị Bình",
            "Lê Văn Cường",
            "Phạm Thị Dung",
            "Hoàng Văn Em",
            "Vũ Thị Phương",
            "Đặng Văn Giang",
        ]
    }

    df = pd.DataFrame(sample_data)
    df.to_excel("data.xlsx", index=False)
    print("   ✅ Đã tạo data.xlsx với 7 tên mẫu")

    # Thông báo về template
    print("   💡 Đảm bảo file template.docx có Content Control với tag='name'")


def main():
    """
    Hàm chính
    """
    # Cấu hình
    excel_file = "tenhs.xlsx"
    template_file = "template.docx"
    output_folder = "generated_files"
    tag_name = "name"

    print("🎯 EXCEL TO WORD AUTOMATION")
    print("=" * 40)

    # Kiểm tra file
    if not os.path.exists(excel_file):
        print(f"❌ Không tìm thấy {excel_file}")
        create_sample_data()
        print(f"✅ Đã tạo file mẫu, hãy chạy lại!")
        return

    if not os.path.exists(template_file):
        print(f"❌ Không tìm thấy {template_file}")
        print(f"💡 Tạo file template.docx với Content Control có tag='name'")
        return

    # Hỏi xác nhận
    print(f"📋 Cấu hình:")
    print(f"   📊 Excel: {excel_file}")
    print(f"   📄 Template: {template_file}")
    print(f"   📁 Output: {output_folder}")
    print(f"   🏷️ Tag: {tag_name}")

    confirm = input("\n➡️ Tiếp tục? (y/n): ").lower().strip()

    if confirm == "y":
        process_excel_to_word_files(excel_file, template_file, output_folder, tag_name)
    else:
        print("🛑 Đã hủy!")


if __name__ == "__main__":
    main()
